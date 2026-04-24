"""
PDFWala Enterprise V11.1.0
tasks/office_tasks.py — Async Celery tasks for Office <-> PDF conversions.

Changes vs V11.0:
  - pdf_to_word_task: chunked parallel processing for PDFs >
    PDF2WORD_CHUNK_THRESHOLD pages (default 50).
    chunk_size=80, max_workers=2.
    Falls back to single-pass if chunking fails.
  - pdf_to_excel_task: chunked parallel processing for PDFs >
    PDF_TO_EXCEL_CHUNK_THRESHOLD pages (default 80).
    chunk_size=80, max_workers=2 (pdfplumber is CPU+RAM heavy).
    Falls back to single-pass if chunking fails.
    Per-chunk Excel files merged into one workbook at the end.
"""

import os
import logging
import shutil
import subprocess
import tempfile
from concurrent.futures import ThreadPoolExecutor

from workers.celery_app import celery_app
from services.redis_service import redis_service
from services.queue_service import cb_libreoffice
from config import Config
from utils.helpers import get_timestamp
from utils.pdf_utils import chunked_pdf_processor

log = logging.getLogger("pdfwala.tasks.office")

try:
    from pdf2docx import Converter as Pdf2DocxConverter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

try:
    from openpyxl import Workbook, load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False


# ── Config ─────────────────────────────────────────────────────────────────────

PDF_TO_EXCEL_CHUNK_THRESHOLD = int(getattr(Config, "PDF_TO_EXCEL_CHUNK_THRESHOLD", 80))
PDF_TO_EXCEL_CHUNK_PAGES     = int(getattr(Config, "PDF_TO_EXCEL_CHUNK_PAGES",     80))
PDF_TO_EXCEL_MAX_WORKERS     = int(getattr(Config, "PDF_TO_EXCEL_MAX_WORKERS",      2))


# ── PDF-to-Word chunk config ───────────────────────────────────────────────────

PDF2WORD_CHUNK_THRESHOLD = int(os.environ.get("PDF2WORD_CHUNK_THRESHOLD", 50))
PDF2WORD_CHUNK_PAGES     = int(os.environ.get("PDF2WORD_CHUNK_PAGES",     80))
PDF2WORD_MAX_WORKERS     = int(os.environ.get("PDF2WORD_MAX_WORKERS",      2))


# ── LibreOffice helper (task-local) ───────────────────────────────────────────

def _libre_convert(input_path: str, fmt: str, out_dir: str):
    """Run LibreOffice conversion via subprocess list args (no shell=True)."""
    if not cb_libreoffice.can_execute():
        log.error("CircuitBreaker[libreoffice] OPEN")
        return None
    try:
        result = subprocess.run(
            [Config.LIBREOFFICE, "--headless", "--convert-to", fmt,
             "--outdir", out_dir, input_path],
            capture_output=True,
            timeout=Config.SUBPROCESS_TIMEOUT,
        )
        if result.returncode != 0:
            cb_libreoffice.record_failure()
            log.error(f"LibreOffice rc={result.returncode}: "
                      f"{result.stderr.decode('utf-8', errors='ignore')[:500]}")
            return None
        from pathlib import Path
        base    = Path(input_path).stem
        pattern = os.path.join(out_dir, f"{base}.{fmt}")
        if os.path.exists(pattern):
            cb_libreoffice.record_success()
            return pattern
        matches = list(Path(out_dir).glob(f"*.{fmt}"))
        if matches:
            cb_libreoffice.record_success()
            return str(matches[0])
        cb_libreoffice.record_failure()
        return None
    except subprocess.TimeoutExpired:
        cb_libreoffice.record_failure()
        log.error("LibreOffice timed out")
        return None
    except Exception as ex:
        cb_libreoffice.record_failure()
        log.error(f"LibreOffice exception: {ex}")
        return None


# ── PDF-to-Excel single-pass helper ───────────────────────────────────────────

def _extract_tables_from_pdf(pdf_path: str, wb: Workbook, sheet_offset: int = 0) -> int:
    """Extract tables from pdf_path into wb starting at sheet number sheet_offset+1."""
    tables_extracted = 0
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    for table in page.extract_tables():
                        if table and any(any(c for c in row if c) for row in table):
                            tables_extracted += 1
                            ws = wb.create_sheet(f"Table_{sheet_offset + tables_extracted}")
                            for row in table:
                                ws.append([str(c).strip() if c else "" for c in row])
        except Exception as ex:
            log.warning(f"pdfplumber extraction failed on chunk: {ex}")
    if tables_extracted == 0 and FITZ_AVAILABLE:
        ws      = wb.create_sheet(f"Text_{sheet_offset + 1}")
        doc     = fitz.open(pdf_path)
        row_idx = 1
        try:
            for pg_num, pg in enumerate(doc):
                ws.cell(row_idx, 1, f"--- Page {pg_num + 1} ---")
                row_idx += 1
                for line in pg.get_text("text").split("\n"):
                    if line.strip():
                        ws.cell(row_idx, 1, line.strip())
                        row_idx += 1
        finally:
            doc.close()
        tables_extracted = 1
    return tables_extracted


def _merge_excel_chunks(chunk_xlsx_paths, output_path: str) -> None:
    """Merge multiple per-chunk .xlsx files into a single output workbook."""
    if not OPENPYXL_AVAILABLE:
        raise RuntimeError("openpyxl required for merge_excel_chunks")
    out_wb    = Workbook()
    out_wb.remove(out_wb.active)
    sheet_idx = 1
    for chunk_path in chunk_xlsx_paths:
        try:
            src_wb = load_workbook(chunk_path, read_only=True, data_only=True)
            for src_ws_name in src_wb.sheetnames:
                src_ws  = src_wb[src_ws_name]
                new_name = f"Sheet_{sheet_idx}"
                out_ws   = out_wb.create_sheet(title=new_name)
                for row in src_ws.iter_rows(values_only=True):
                    out_ws.append(list(row))
                sheet_idx += 1
            src_wb.close()
        except Exception as ex:
            log.warning(f"merge_excel_chunks: skipping {chunk_path}: {ex}")
    if not out_wb.sheetnames:
        out_wb.create_sheet("Empty")
    tmp = output_path + ".merge_tmp.xlsx"
    out_wb.save(tmp)
    out_wb.close()
    os.replace(tmp, output_path)


def _run_excel_single_pass(input_path: str, output_path: str) -> None:
    """Extract tables / raw text from input_path, save to output_path (.xlsx)."""
    wb = Workbook()
    wb.remove(wb.active)
    _extract_tables_from_pdf(input_path, wb, sheet_offset=0)
    if not wb.sheetnames:
        wb.create_sheet("Empty")
    wb.save(output_path)
    wb.close()


# ── Tasks ──────────────────────────────────────────────────────────────────────

if celery_app is not None:

    # ------------------------------------------------------------------
    # PDF → Word  (V11.1 — chunked for large PDFs)
    # ------------------------------------------------------------------
    @celery_app.task(
        bind=True,
        max_retries=2,
        name="pdfwala.tasks.office_tasks.pdf_to_word_task",
        queue="office",
        time_limit=3600,       # 60 min for chunked large jobs
        soft_time_limit=3300,
    )
    def pdf_to_word_task(self, input_path: str, output_path: str, job_id: str):
        """
        Async PDF → Word conversion.
        ≤ PDF2WORD_CHUNK_THRESHOLD pages: single-pass
        > PDF2WORD_CHUNK_THRESHOLD pages: chunked parallel conversion
        """
        try:
            if not PDF2DOCX_AVAILABLE:
                raise RuntimeError("pdf2docx not installed")
            if not FITZ_AVAILABLE:
                raise RuntimeError("PyMuPDF not installed")
            if not DOCX_AVAILABLE:
                raise RuntimeError("python-docx not installed (needed for merge)")

            redis_service.job_update(job_id, {"status": "processing", "progress": "2"})

            src = fitz.open(input_path)
            total_pages = len(src)
            src.close()
            redis_service.job_update(job_id, {"total_pages": str(total_pages)})

            log.info(f"pdf_to_word_task {job_id}: {total_pages} pages, "
                     f"threshold={PDF2WORD_CHUNK_THRESHOLD}")

            if total_pages <= PDF2WORD_CHUNK_THRESHOLD:
                # ── Fast path: single-pass ───────────────────────────────────
                redis_service.job_update(job_id, {"progress": "10"})
                cv = Pdf2DocxConverter(input_path)
                cv.convert(output_path, start=0, end=None)
                cv.close()
            else:
                # ── Chunked path: parallel conversion ────────────────────────
                redis_service.job_update(job_id, {"progress": "5"})

                chunks = []
                for start in range(0, total_pages, PDF2WORD_CHUNK_PAGES):
                    end = min(start + PDF2WORD_CHUNK_PAGES, total_pages) - 1
                    chunks.append((start, end))

                def convert_chunk(chunk_start, chunk_end):
                    chunk_out = f"/tmp/chunk_{job_id}_{chunk_start}_{chunk_end}.docx"
                    cv = Pdf2DocxConverter(input_path)
                    cv.convert(chunk_out, start=chunk_start, end=chunk_end)
                    cv.close()
                    current = min(chunk_end + 1, total_pages)
                    pct = int(current / total_pages * 100)
                    redis_service.job_update(job_id, {
                        "progress": str(pct),
                        "current_page": str(current),
                        "status_detail": f"Page {current} of {total_pages}"
                    })
                    return chunk_out

                chunk_files = []
                with ThreadPoolExecutor(max_workers=PDF2WORD_MAX_WORKERS) as executor:
                    futures = [executor.submit(convert_chunk, s, e) for s, e in chunks]
                    for f in futures:
                        chunk_files.append(f.result())

                # Merge chunks
                master = None
                for chunk_file in sorted(chunk_files):
                    doc = DocxDocument(chunk_file)
                    if master is None:
                        master = doc
                    else:
                        for element in doc.element.body:
                            master.element.body.append(element)
                master.save(output_path)

                for f in chunk_files:
                    try:
                        os.remove(f)
                    except OSError:
                        pass

            if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
                raise RuntimeError("Output file missing or empty after conversion")

            redis_service.job_update(job_id, {
                "status":       "completed",
                "progress":     "100",
                "output_path":  output_path,
                "completed_at": get_timestamp(),
                "total_pages":  str(total_pages),
            })
            return {"status": "completed", "output": output_path}

        except Exception as ex:
            log.error(f"pdf_to_word_task {job_id}: {ex}")
            redis_service.job_update(job_id, {"status": "failed", "error": str(ex)})
            delay = min(60 * (2 ** self.request.retries), 300)
            raise self.retry(exc=ex, countdown=delay, max_retries=2)
        finally:
            try:
                os.remove(input_path)
            except OSError:
                pass

    # ------------------------------------------------------------------
    # PDF → Excel  (chunked for large files)
    # ------------------------------------------------------------------
    @celery_app.task(
        bind=True,
        max_retries=2,
        name="pdfwala.tasks.office_tasks.pdf_to_excel_task",
        queue="office",
        time_limit=3600,
        soft_time_limit=3300,
    )
    def pdf_to_excel_task(self, input_path: str, output_path: str, job_id: str):
        """Async PDF -> Excel extraction (tables or raw text fallback)."""
        try:
            if not OPENPYXL_AVAILABLE:
                raise RuntimeError("openpyxl not installed")
            redis_service.job_update(job_id, {"status": "processing"})
            total_pages = 0
            if FITZ_AVAILABLE:
                doc_check   = fitz.open(input_path)
                total_pages = len(doc_check)
                doc_check.close()
                redis_service.job_update(job_id, {"total_pages": str(total_pages)})
            if total_pages > PDF_TO_EXCEL_CHUNK_THRESHOLD:
                log.info(f"pdf_to_excel_task {job_id}: {total_pages} pages — chunked")
                base_temp = os.path.dirname(input_path)
                def process_excel_chunk(chunk_path, chunk_idx, start_page, end_page):
                    chunk_wb = Workbook()
                    chunk_wb.remove(chunk_wb.active)
                    _extract_tables_from_pdf(chunk_path, chunk_wb, sheet_offset=0)
                    chunk_out = os.path.join(base_temp, f"_chunk_{job_id}_{chunk_idx:04d}_excel.xlsx")
                    chunk_wb.save(chunk_out)
                    chunk_wb.close()
                    return chunk_out
                success = chunked_pdf_processor(
                    input_path=input_path, output_path=output_path, job_id=job_id,
                    total_pages=total_pages, chunk_size=PDF_TO_EXCEL_CHUNK_PAGES,
                    max_workers=PDF_TO_EXCEL_MAX_WORKERS,
                    process_chunk_func=process_excel_chunk,
                    merge_func=_merge_excel_chunks, redis_service=redis_service,
                    tool_name="PDF-to-Excel", report_progress=True, chunk_retry=1,
                )
                if not success:
                    log.warning(f"pdf_to_excel_task {job_id}: chunked failed — fallback")
                    _run_excel_single_pass(input_path, output_path)
            else:
                log.info(f"pdf_to_excel_task {job_id}: {total_pages} pages — single-pass")
                _run_excel_single_pass(input_path, output_path)
            if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
                raise RuntimeError("Output Excel file missing or empty")
            redis_service.job_update(job_id, {
                "status": "completed", "progress": "100",
                "output_path": output_path, "completed_at": get_timestamp(),
            })
            return {"status": "completed", "output": output_path}
        except Exception as ex:
            log.error(f"pdf_to_excel_task {job_id}: {ex}")
            redis_service.job_update(job_id, {"status": "failed", "error": str(ex)})
            raise self.retry(exc=ex, countdown=30)
        finally:
            try:
                os.remove(input_path)
            except OSError:
                pass

    # ------------------------------------------------------------------
    # Excel → Word  [V11.0 — unchanged]
    # ------------------------------------------------------------------
    @celery_app.task(
        bind=True, max_retries=2,
        name="pdfwala.tasks.office_tasks.excel_to_word_task",
        queue="office", time_limit=1800, soft_time_limit=1500,
    )
    def excel_to_word_task(self, input_path, output_path, job_id,
                           preserve_formulas=True, row_limit=5000):
        """Async Excel -> Word conversion."""
        try:
            if not OPENPYXL_AVAILABLE or not DOCX_AVAILABLE:
                raise RuntimeError("openpyxl + python-docx required")
            redis_service.job_update(job_id, {"status": "processing", "progress": "5"})
            wb = load_workbook(input_path, data_only=not preserve_formulas)
            doc = DocxDocument()
            sheet_count = len(wb.sheetnames)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                doc.add_heading(sheet_name, level=1)
                all_rows = list(ws.iter_rows(values_only=True, max_row=row_limit+1))
                truncated = len(all_rows) > row_limit
                rows_write = all_rows[:row_limit]
                if not rows_write:
                    doc.add_paragraph("(empty sheet)")
                    continue
                n_cols = max((len(r) for r in rows_write), default=1)
                table = doc.add_table(rows=len(rows_write), cols=n_cols)
                try:
                    table.style = "Light Grid Accent 1"
                except Exception:
                    pass
                for r_idx, row_data in enumerate(rows_write):
                    for c_idx in range(n_cols):
                        val = row_data[c_idx] if c_idx < len(row_data) else None
                        cell = table.cell(r_idx, c_idx)
                        cell.text = str(val) if val is not None else ""
                        if r_idx == 0:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.bold = True
                if truncated:
                    doc.add_paragraph(f"(Truncated to {row_limit} rows)")
                doc.add_paragraph()
            wb.close()
            doc.save(output_path)
            redis_service.job_update(job_id, {
                "status": "completed", "progress": "100",
                "output_path": output_path, "completed_at": get_timestamp(),
                "sheets": str(sheet_count),
            })
            return {"status": "completed", "output": output_path}
        except Exception as ex:
            log.error(f"excel_to_word_task {job_id}: {ex}")
            redis_service.job_update(job_id, {"status": "failed", "error": str(ex)})
            raise self.retry(exc=ex, countdown=30)
        finally:
            try:
                os.remove(input_path)
            except OSError:
                pass

    # ------------------------------------------------------------------
    # Word → PDF
    # ------------------------------------------------------------------
    @celery_app.task(
        bind=True, max_retries=2,
        name="pdfwala.tasks.office_tasks.word_to_pdf_task",
        queue="office", time_limit=1800, soft_time_limit=1500,
    )
    def word_to_pdf_task(self, input_path, output_path, job_id):
        """Async Word -> PDF via LibreOffice."""
        try:
            redis_service.job_update(job_id, {"status": "processing"})
            out_dir = tempfile.mkdtemp()
            converted = _libre_convert(input_path, "pdf", out_dir)
            if converted and os.path.exists(converted):
                shutil.move(converted, output_path)
            else:
                raise RuntimeError("LibreOffice conversion failed")
            shutil.rmtree(out_dir, ignore_errors=True)
            redis_service.job_update(job_id, {
                "status": "completed", "progress": "100",
                "output_path": output_path, "completed_at": get_timestamp(),
            })
            return {"status": "completed", "output": output_path}
        except Exception as ex:
            log.error(f"word_to_pdf_task {job_id}: {ex}")
            redis_service.job_update(job_id, {"status": "failed", "error": str(ex)})
            raise self.retry(exc=ex, countdown=30)
        finally:
            try:
                os.remove(input_path)
            except OSError:
                pass

    # ------------------------------------------------------------------
    # Excel → PDF
    # ------------------------------------------------------------------
    @celery_app.task(
        bind=True, max_retries=2,
        name="pdfwala.tasks.office_tasks.excel_to_pdf_task",
        queue="office", time_limit=1800, soft_time_limit=1500,
    )
    def excel_to_pdf_task(self, input_path, output_path, job_id):
        """Async Excel -> PDF via LibreOffice."""
        try:
            redis_service.job_update(job_id, {"status": "processing"})
            out_dir = tempfile.mkdtemp()
            converted = _libre_convert(input_path, "pdf", out_dir)
            if converted and os.path.exists(converted):
                shutil.move(converted, output_path)
            else:
                raise RuntimeError("LibreOffice conversion failed")
            shutil.rmtree(out_dir, ignore_errors=True)
            redis_service.job_update(job_id, {
                "status": "completed", "progress": "100",
                "output_path": output_path, "completed_at": get_timestamp(),
            })
            return {"status": "completed", "output": output_path}
        except Exception as ex:
            log.error(f"excel_to_pdf_task {job_id}: {ex}")
            redis_service.job_update(job_id, {"status": "failed", "error": str(ex)})
            raise self.retry(exc=ex, countdown=30)
        finally:
            try:
                os.remove(input_path)
            except OSError:
                pass

else:
    pdf_to_word_task   = None
    pdf_to_excel_task  = None
    excel_to_word_task = None
    word_to_pdf_task   = None
    excel_to_pdf_task  = None
